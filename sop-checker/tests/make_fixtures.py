"""Fixture generator.

Everything here builds .docx files with python-docx so the fixtures rebuild on
demand -- no binary blobs checked in. Low-level OOXML builders (fields, theme
fonts, split runs, text boxes, nested tables, multi-section footers) are reused
by both the extraction tests and the per-rule fixtures.

Run directly to (re)write every fixture to tests/fixtures/:

    python -m tests.make_fixtures
"""
from __future__ import annotations

import io
import os
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# make sure the VML prefix used for text boxes resolves in OxmlElement
nsmap.setdefault("v", "urn:schemas-microsoft-com:vml")

FIXTURE_DIR = os.path.join(os.path.dirname(__file__), "fixtures")


# --------------------------------------------------------------------------
# low-level OOXML builders
# --------------------------------------------------------------------------
def _rpr(r):
    rpr = r.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        r.insert(0, rpr)
    return rpr


def raw_run(text, *, ascii=None, theme=None, size=None, bold=None,
            italic=None, color=None, underline=None):
    """Build a bare <w:r> with explicit formatting (not tied to a Run obj)."""
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    r.append(rpr)
    if ascii or theme:
        rf = OxmlElement("w:rFonts")
        if ascii:
            rf.set(qn("w:ascii"), ascii)
            rf.set(qn("w:hAnsi"), ascii)
        if theme:
            rf.set(qn("w:asciiTheme"), theme)
            rf.set(qn("w:hAnsiTheme"), theme)
        rpr.append(rf)
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rpr.append(sz)
    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), underline)
        rpr.append(u)
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rpr.append(c)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def add_raw_run(paragraph, text, **kw):
    r = raw_run(text, **kw)
    paragraph._p.append(r)
    return r


def add_complex_field(paragraph, instr, result="1"):
    """Insert a complex field (fldChar begin/instr/separate/result/end)."""
    p = paragraph._p

    def wrap(child):
        r = OxmlElement("w:r")
        r.append(child)
        return r

    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate")
    rt = OxmlElement("w:t")
    rt.text = result
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    for child in (fb, it, fs):
        p.append(wrap(child))
    rr = OxmlElement("w:r")
    rr.append(rt)
    p.append(rr)
    p.append(wrap(fe))


def add_simple_field(paragraph, instr, result="1"):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = result
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def add_split_word(paragraph, fragments, *, ascii=None, theme=None, size=None):
    """Append a word split across several runs, mimicking spell-check splits,
    with inert w:proofErr markers between them."""
    p = paragraph._p
    start = OxmlElement("w:proofErr")
    start.set(qn("w:type"), "spellStart")
    p.append(start)
    for frag in fragments:
        p.append(raw_run(frag, ascii=ascii, theme=theme, size=size))
    end = OxmlElement("w:proofErr")
    end.set(qn("w:type"), "spellEnd")
    p.append(end)


def add_textbox(paragraph, lines):
    """Attach a VML text box (w:pict/v:shape/v:textbox/w:txbxContent) to a run
    inside ``paragraph``. ``lines`` become paragraphs inside the box."""
    r = OxmlElement("w:r")
    pict = OxmlElement("w:pict")
    shape = OxmlElement("v:shape")
    shape.set("style", "width:200pt;height:50pt")
    textbox = OxmlElement("v:textbox")
    txbx = OxmlElement("w:txbxContent")
    for line in lines:
        tp = OxmlElement("w:p")
        tr = raw_run(line)
        tp.append(tr)
        txbx.append(tp)
    textbox.append(txbx)
    shape.append(textbox)
    pict.append(shape)
    r.append(pict)
    paragraph._p.append(r)


def add_nested_table(cell, rows, cols, filler="x"):
    """Add a table inside a table cell and return it."""
    tbl = cell.add_table(rows, cols) if hasattr(cell, "add_table") else None
    if tbl is None:
        # python-docx _Cell.add_table takes (rows, cols)
        tbl = cell.add_table(rows=rows, cols=cols)
    for r in range(rows):
        for c in range(cols):
            tbl.cell(r, c).text = f"{filler}{r}{c}"
    return tbl


def set_core(doc, *, title=None, author=None, subject=None, keywords=None,
             last_modified_by=None, created=None, modified=None):
    cp = doc.core_properties
    if title is not None:
        cp.title = title
    if author is not None:
        cp.author = author
    if subject is not None:
        cp.subject = subject
    if keywords is not None:
        cp.keywords = keywords
    if last_modified_by is not None:
        cp.last_modified_by = last_modified_by
    if created is not None:
        cp.created = created
    if modified is not None:
        cp.modified = modified


def body_para(doc, text, *, font="Calibri", size=11):
    """A plain body paragraph with an explicit literal font/size."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    return p


def to_bytes(doc) -> io.BytesIO:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# --------------------------------------------------------------------------
# extraction edge-case documents (returned as Document objects)
# --------------------------------------------------------------------------
def doc_interleave():
    """Paragraph, table, paragraph, table, paragraph -> flow order."""
    d = Document()
    d.add_paragraph("Alpha")
    t1 = d.add_table(rows=1, cols=1)
    t1.cell(0, 0).text = "Table-one-cell"
    d.add_paragraph("Beta")
    t2 = d.add_table(rows=1, cols=1)
    t2.cell(0, 0).text = "Table-two-cell"
    d.add_paragraph("Gamma")
    return d


def doc_nested_table():
    d = Document()
    d.add_paragraph("Before")
    outer = d.add_table(rows=1, cols=1)
    cell = outer.cell(0, 0)
    cell.paragraphs[0].text = "OuterCell"
    add_nested_table(cell, 2, 2, filler="n")
    d.add_paragraph("After")
    return d


def doc_split_word():
    d = Document()
    p = d.add_paragraph()
    add_split_word(p, ["proof", "read", "ing"], ascii="Arial", size=11)
    return d


def doc_theme_vs_literal():
    """Two runs, one theme minor + one literal of the same face."""
    d = Document()
    # discover the template's minor face so the literal run matches it
    p = d.add_paragraph()
    add_raw_run(p, "themed ", theme="minorHAnsi")
    # literal run set to the same resolved face is added by the test after it
    # reads doc.styles.theme_minor; here we also add a Cambria literal that
    # matches the default template's minor font.
    add_raw_run(p, "literal", ascii="Cambria")
    return d


def doc_textbox():
    d = Document()
    p = d.add_paragraph("Anchor paragraph")
    add_textbox(p, ["Text box line one", "Text box line two"])
    d.add_paragraph("Following paragraph")
    return d


def doc_three_sections_diff_first():
    d = Document()
    d.add_paragraph("Section one body")
    for i in range(2):
        sec = d.add_section()
        d.add_paragraph(f"Section {i + 2} body")
    for i, sec in enumerate(d.sections):
        sec.different_first_page_header_footer = True
        add_simple_field(sec.first_page_footer.paragraphs[0], "PAGE")
        add_complex_field(sec.footer.paragraphs[0], " PAGE ")
    return d


def doc_no_headers_footers():
    d = Document()
    d.add_paragraph("Body only, no headers or footers.")
    # ensure no footer/header content is added
    return d


def doc_both_field_kinds():
    d = Document()
    p1 = d.add_paragraph("Simple: ")
    add_simple_field(p1, "PAGE \\* MERGEFORMAT")
    p2 = d.add_paragraph("Complex: ")
    add_complex_field(p2, " PAGE ")
    p3 = d.add_paragraph("Pages: ")
    add_complex_field(p3, " NUMPAGES ")
    return d


SIMPLE_PROSE = (
    "This guide shows how we check our work. We want each step to be clear. "
    "The team reads this guide before they start a new job. Each person "
    "signs the form when the work is done. We keep the forms in a safe place. "
    "If a step is not clear, ask the team lead for help. The lead will show "
    "you what to do. We look at this guide once a year. We fix any part that "
    "is out of date. Good notes help us learn and get better each day."
)
DENSE_PROSE = (
    "Notwithstanding the aforementioned considerations, the comprehensive "
    "operationalisation of multifaceted organisational methodologies "
    "necessitates the meticulous synchronisation of heterogeneous "
    "infrastructural dependencies, thereby engendering substantial "
    "epistemological ramifications whose intricate interdependencies "
    "invariably confound stakeholders endeavouring to disambiguate the "
    "underlying procedural abstractions notwithstanding exhaustive "
    "documentation efforts undertaken throughout preceding developmental "
    "iterations and subsequent retrospective evaluations thereof."
)


def _footer_with_fields(doc, *, doc_id=True, confidentiality=True,
                        page_field=True, hardcoded=False):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    prefix = []
    if doc_id:
        prefix.append("QMS-SOP-001")
    if confidentiality:
        prefix.append("Confidential")
    p.text = "  |  ".join(prefix)
    if prefix:
        p.add_run("  |  Page ")
    else:
        p.add_run("Page ")
    if hardcoded:
        p.add_run("1")
    elif page_field:
        add_complex_field(p, " PAGE ")
        p.add_run(" of ")
        add_complex_field(p, " NUMPAGES ", result="5")


def _revision_table(doc, *, dates, author="Jane Smith"):
    tbl = doc.add_table(rows=1 + len(dates), cols=4)
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text = "Version", "Date"
    hdr[2].text, hdr[3].text = "Author", "Description"
    versions = ["1.0", "1.1", "1.2", "1.3", "1.4"]
    for i, d in enumerate(dates):
        row = tbl.rows[i + 1].cells
        row[0].text = versions[i] if i < len(versions) else "1.x"
        row[1].text = d
        row[2].text = author
        row[3].text = "Update"
    return tbl


def unstyled_sop(*, doc_info=True, sig_date="25/08/2026"):
    """An SOP written the way plenty of real ones are: not one heading style
    in the file, sections marked only by a bigger bold line, and the title,
    version and author stated in a document-information table rather than in
    metadata. The core author is the generating library, as it is whenever a
    file is produced by a script.

    This is the shape that used to leave rules 1, 4 and 12 unevaluated.
    """
    d = Document()
    set_core(d, author="python-docx")

    def visual(text, size, *, bold=True, center=False):
        para = d.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Cambria"
        run.font.size = Pt(size)
        run.bold = bold
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return para

    visual("Document Reviewer - Test Document", 18, center=True)
    visual("Sample document for testing selected compliance rules", 10,
           bold=False, center=True)

    if doc_info:
        info = [("Document Information", "Value"),
                ("Document Title", "Employee Onboarding SOP"),
                ("Document ID", "SOP-001"),
                ("Version", "2.1"),
                ("Author", "John Smith"),
                ("Role", "Quality Manager")]
        tbl = d.add_table(rows=len(info), cols=2)
        for row, (label, value) in zip(tbl.rows, info):
            row.cells[0].text, row.cells[1].text = label, value

    d.add_paragraph()
    for heading, body in [
        ("Objective", "The objective of this procedure is to provide a "
                      "consistent approach for onboarding new employees "
                      "across every department of the organisation."),
        ("Scope", "This procedure applies to all new employees joining the "
                  "organisation, including contractors and temporary staff "
                  "engaged for longer than one month."),
        ("Responsibilities", "The Quality Manager is responsible for "
                             "ensuring that the onboarding records are "
                             "complete, accurate and retained for audit."),
    ]:
        visual(heading, 14)
        body_para(d, body, font="Cambria")

    visual("Revision History", 14)
    tbl = d.add_table(rows=4, cols=3)
    rows = [("Version", "Date", "Description"),
            ("1.0", "12/01/2025", "Initial release"),
            ("2.0", "20/06/2025", "Major procedure update"),
            ("2.1", "15/08/2026", "Minor update")]
    for row, values in zip(tbl.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    # the approval block sits directly under the revision table, so the
    # table's last date is as close to it as its own
    visual("Approval", 14)
    body_para(d, "Approved By: John Smith", font="Cambria")
    body_para(d, "Signature: ______________________________", font="Cambria")
    body_para(d, f"Date: {sig_date}", font="Cambria")

    visual("Procedure", 14)
    body_para(d, "The manager reviews the employee information, confirms the "
                 "required documents are on file, and records the outcome in "
                 "the onboarding register before the end of the first week.",
              font="Cambria")

    _footer_with_fields(d, page_field=False)
    return d


def golden_sop(*, break_author=False, break_rev_dates=False,
               break_version=False, break_revision_section=False,
               break_signature=False, break_sig_dates=False,
               break_fonts=False, break_required=False,
               break_page_field=False, break_footer_details=False,
               break_readability=False):
    """A complete SOP that passes all 13 rules; keyword flags inject a single
    defect each so per-rule fail fixtures reuse one builder."""
    d = Document()
    author = "User" if break_author else "Jane Smith"
    set_core(
        d,
        title="Quality Control Procedure",
        author=author,
        last_modified_by="Windows User" if break_author else "Jane Smith",
        subject="Version 1.2",
        keywords="quality, control",
    )

    # header states the current version
    header = d.sections[0].header
    header.paragraphs[0].text = (
        "Version 1.1" if break_version else "Version 1.2")

    # ---- title page ----
    d.add_heading("Quality Control Procedure", level=1)
    body_para(d, "Document Version: 1.2")
    if not break_author:
        body_para(d, "Author: Jane Smith, Quality Manager")

    # ---- signatures with inline dates ----
    if not break_signature:
        labels = ["Prepared by", "Reviewed by", "Approved by"]
        # break_author anonymises the signatures too, else a named signer
        # would legitimately satisfy rule 2
        names = ["", "", ""] if break_author \
            else ["Jane Smith", "Raj Patel", "Mary Lee"]
        for label, name in zip(labels, names):
            if name:
                line = f"{label}: {name}"
                if not break_sig_dates:
                    line += ". Date: 01/02/2024"
                body_para(d, line)
            else:
                body_para(d, f"{label}:")
                if not break_sig_dates:
                    body_para(d, "Signed on Date: 01/02/2024")

    # ---- revision history ----
    if not break_revision_section:
        d.add_heading("Revision History", level=1)
        if break_rev_dates:
            dates = ["01/01/2023", "15/06/2022", "03/02/2099"]  # unordered+future
        else:
            dates = ["01/01/2023", "15/06/2023", "03/02/2024"]
        _revision_table(d, dates=dates,
                        author=("User" if break_author else "Jane Smith"))

    # ---- required sections ----
    sections = ["Purpose", "Scope", "Responsibilities", "Procedure",
                "References"]
    if break_required:
        sections.remove("References")
    for name in sections:
        d.add_heading(name, level=1)
        if name == "Procedure":
            prose = DENSE_PROSE if break_readability else SIMPLE_PROSE
            body_para(d, prose)
        else:
            body_para(d, f"This section covers the {name.lower()} of the work.")

    if break_fonts:
        body_para(d, "This whole paragraph uses a different typeface entirely.",
                  font="Times New Roman")

    # ---- footer ----
    _footer_with_fields(
        d,
        doc_id=not break_footer_details,
        confidentiality=not break_footer_details,
        page_field=not break_page_field,
        hardcoded=break_page_field,
    )
    return d


# --------------------------------------------------------------------------
# rule fixtures registered for on-disk generation (phase 3)
# --------------------------------------------------------------------------
RULE_FIXTURES: dict[str, "callable"] = {}


def _register_rule_fixtures():
    RULE_FIXTURES["rule_pass_all"] = lambda: golden_sop()
    flags = [
        ("rule_02_fail", dict(break_author=True)),
        ("rule_03_fail", dict(break_rev_dates=True)),
        ("rule_04_fail", dict(break_version=True)),
        ("rule_05_fail", dict(break_revision_section=True)),
        ("rule_06_fail", dict(break_signature=True)),
        ("rule_07_fail", dict(break_sig_dates=True)),
        ("rule_08_fail", dict(break_fonts=True)),
        ("rule_10_fail", dict(break_required=True)),
        ("rule_11_fail", dict(break_page_field=True)),
        ("rule_12_fail", dict(break_readability=True)),
        ("rule_13_fail", dict(break_footer_details=True)),
    ]
    for name, kw in flags:
        RULE_FIXTURES[name] = (lambda k=kw: golden_sop(**k))


_register_rule_fixtures()


def build_all(outdir: str = FIXTURE_DIR) -> list[str]:
    os.makedirs(outdir, exist_ok=True)
    written = []
    for name, fn in RULE_FIXTURES.items():
        doc = fn()
        path = os.path.join(outdir, f"{name}.docx")
        doc.save(path)
        written.append(path)
    return written


if __name__ == "__main__":
    paths = build_all()
    print(f"wrote {len(paths)} fixtures to {FIXTURE_DIR}")
    for p in paths:
        print("  ", os.path.basename(p))
