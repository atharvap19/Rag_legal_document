"""Single-file .docx extractor: OOXML in, :class:`Doc` model out.

Everything the 13 rules need is produced here -- flow-ordered blocks, resolved
fonts, merged runs, field codes, tables, and per-section headers/footers. The
rules consume only the dataclasses below, never python-docx or lxml, which is
what keeps them unit-testable.

Collapsed from the former extract/ package; the section banners mark what were
previously separate modules.
"""
from __future__ import annotations

from collections import Counter
from lxml import etree
from dataclasses import dataclass, field
from datetime import datetime
from typing import Iterable, Optional, Union
import re
from docx import Document as OpenDocument


# ========================================================================
# OOXML namespace helpers
# ========================================================================
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
#http://schemas.openxmlformats.org/wordprocessingml/2006/main — Handles core document content 
#like paragraphs (w:p), runs (w:r), and text (w:t).
A = "http://schemas.openxmlformats.org/drawingml/2006/main"

WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
V = "urn:schemas-microsoft-com:vml"

NS = {"w": W, "a": A, "wp": WP, "r": R, "v": V}


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def a(tag: str) -> str:
    return f"{{{A}}}{tag}"


def local(el: etree._Element) -> str:
    """Local tag name without namespace."""
    tag = el.tag
    if isinstance(tag, str) and "}" in tag:
        return tag.rsplit("}", 1)[1]
    return tag if isinstance(tag, str) else ""


def wval(el: Optional[etree._Element], default=None) -> Optional[str]:
    """Return the ``w:val`` attribute of an element, or default."""
    if el is None:
        return default
    v = el.get(w("val"))
    return v if v is not None else default


def get(parent: Optional[etree._Element], *tags: str) -> Optional[etree._Element]:
    """Chase a path of ``w:`` child tags, returning the leaf or None."""
    cur = parent
    for tag in tags:
        if cur is None:
            return None
        cur = cur.find(w(tag))
    return cur


def as_bool(el: Optional[etree._Element]) -> Optional[bool]:
    """Interpret an on/off toggle element (w:b, w:i, ...).

    Present with no val, or val in {1,true,on} => True.
    val in {0,false,off} => False. Absent => None.
    """
    if el is None:
        return None
    v = el.get(w("val"))
    if v is None:
        return True
    return v.lower() in ("1", "true", "on")


def twips_to_pt(value: Optional[str]) -> Optional[float]:
    if value is None:
        return None
    try:
        return int(value) / 20.0
    except (TypeError, ValueError):
        return None


def halfpt_to_pt(value: Optional[str]) -> Optional[float]:
    if value is None:
        return None
    try:
        return int(value) / 2.0
    except (TypeError, ValueError):
        return None


# ========================================================================
# The Doc model -- what the rules consume
# ========================================================================
# --------------------------------------------------------------------------
# Run-level
# --------------------------------------------------------------------------
@dataclass
class ResolvedFont:
    """Fully-resolved run formatting (after walking the style chain)."""

    name: Optional[str] = None          # resolved typeface, theme-expanded
    size: Optional[float] = None        # points
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[str] = None     # e.g. "single", "none", or None
    color: Optional[str] = None         # "RRGGBB" hex, "auto", or None
    name_is_theme: bool = False         # name came from a theme font ref

    def key(self) -> tuple:
        """Identity used for run merging and formatting comparison.

        """
        return (self.name, self.size, self.bold, self.italic,
                self.underline, self.color)


@dataclass
class Run:
    text: str
    font: ResolvedFont


@dataclass
class Field:
    """A field code, whether complex (w:fldChar) or simple (w:fldSimple)."""

    kind: str            # first instruction token upper-cased: PAGE, NUMPAGES...
    instruction: str     # full instruction text
    result: str          # cached result text (may be empty)
    simple: bool         # True => w:fldSimple


# --------------------------------------------------------------------------
# Paragraph-level
# --------------------------------------------------------------------------
@dataclass
class ParagraphProps:
    #how is the paragraph formatted 
    style_id: Optional[str] = None
    style_name: Optional[str] = None
    outline_level: Optional[int] = None      # 0-based; None = body text
    heading_level: Optional[int] = None      # 1..9 or None
    space_before: Optional[float] = None     # points
    space_after: Optional[float] = None      # points
    line_spacing: Optional[float] = None     # multiple (auto) or points
    line_spacing_rule: Optional[str] = None  # AUTO | EXACT | AT_LEAST
    alignment: Optional[str] = None          # LEFT | CENTER | RIGHT | JUSTIFY
    first_line_indent: Optional[float] = None  # points (negative => hanging)
    left_indent: Optional[float] = None
    right_indent: Optional[float] = None
    is_list: bool = False
    in_table: bool = False
    # heuristic, never a Word declaration: set only for documents that used
    # direct formatting instead of heading styles (see infer_heading_levels)
    inferred_heading_level: Optional[int] = None


@dataclass
class Paragraph:
    # what the paragraph contains and how it is formatted
    block_index: int
    text: str
    runs: list[Run]
    props: ParagraphProps
    fields: list[Field] = field(default_factory=list)
    in_table: bool = False
    from_textbox: bool = False
    # (table_index, row, col) when this paragraph lives in a table cell
    table_pos: Optional[tuple[int, int, int]] = None
    location: str = ""
    kind: str = "paragraph"

    def word_count(self) -> int:
        #word count 
        return len(self.text.split())


@dataclass
class TableRef:
    """Marker occupying a table's slot in the top-level flow."""

    block_index: int
    table_index: int
    location: str = ""
    kind: str = "table"


Block = Union[Paragraph, TableRef]


# --------------------------------------------------------------------------
# Tables
# --------------------------------------------------------------------------
@dataclass
class Cell:
    blocks: list[Block] = field(default_factory=list)

    def paragraphs(self) -> list[Paragraph]:
        return [b for b in self.blocks if isinstance(b, Paragraph)]

    def text(self) -> str:
        return "\n".join(p.text for p in self.paragraphs())


@dataclass
class Row:
    cells: list[Cell] = field(default_factory=list)


@dataclass
class Table:
    table_index: int
    block_index: int
    rows: list[Row] = field(default_factory=list)

    def iter_paragraphs(self) -> Iterable[Paragraph]:
        for row in self.rows:
            for cell in row.cells:
                for b in cell.blocks:
                    if isinstance(b, Paragraph):
                        yield b


# --------------------------------------------------------------------------
# Sections / headers / footers
# --------------------------------------------------------------------------
@dataclass
class HeaderFooter:
    which: str               # "default" | "first" | "even"
    kind: str                # "header" | "footer"
    section_index: int
    paragraphs: list[Paragraph] = field(default_factory=list)
    fields: list[Field] = field(default_factory=list)
    is_linked_to_previous: bool = False

    def text(self) -> str:
        return "\n".join(p.text for p in self.paragraphs)

    @property
    def location(self) -> str:
        label = self.kind.capitalize()
        if self.which != "default":
            label = f"{self.which.capitalize()}-page {self.kind}"
        return f"{label}, section {self.section_index + 1}"


@dataclass
class Section:
    index: int
    different_first_page: bool = False
    headers: dict[str, HeaderFooter] = field(default_factory=dict)
    footers: dict[str, HeaderFooter] = field(default_factory=dict)

    def all_headers(self) -> list[HeaderFooter]:
        return list(self.headers.values())

    def all_footers(self) -> list[HeaderFooter]:
        return list(self.footers.values())


# --------------------------------------------------------------------------
# Document-level
# --------------------------------------------------------------------------
@dataclass
class CoreProps:
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    keywords: Optional[str] = None
    created: Optional[datetime] = None
    modified: Optional[datetime] = None
    last_modified_by: Optional[str] = None


@dataclass
class HeadingEntry:
    level: int          # heading level (1..9)
    text: str
    block_index: int
    location: str


@dataclass
class StyleInfo:
    style_id: str
    name: Optional[str]
    type: str                       # paragraph | character | table | numbering
    based_on: Optional[str] = None
    is_default: bool = False


@dataclass
class StyleIndex:
    """Read-only summary of the style tree, handed to the rules."""

    styles: dict[str, StyleInfo] = field(default_factory=dict)
    theme_major: Optional[str] = None
    theme_minor: Optional[str] = None
    default_para_style: Optional[str] = None
    default_char_style: Optional[str] = None


@dataclass
class Doc:
    filename: str
    core: CoreProps
    blocks: list[Block]            # top-level flow order
    tables: list[Table]
    sections: list[Section]
    styles: StyleIndex

    # ------------------------------------------------------------------
    # Flow-order helpers (rule 7 leans on these)
    # ------------------------------------------------------------------
    def iter_paragraphs(self) -> Iterable[Paragraph]:
        """Every body paragraph in true flow order (incl. table cells,
        text boxes, nested tables). Headers/footers are *not* included."""
        yield from _iter_para(self.blocks, self.tables)

    def body_paragraphs(self) -> list[Paragraph]:
        return list(self.iter_paragraphs())

    def flow_ordered(self) -> list[Paragraph]:
        """Paragraphs sorted by their stable block_index."""
        return sorted(self.iter_paragraphs(), key=lambda p: p.block_index)

    def all_headers(self) -> list[HeaderFooter]:
        out: list[HeaderFooter] = []
        for s in self.sections:
            out.extend(s.all_headers())
        return out

    def all_footers(self) -> list[HeaderFooter]:
        out: list[HeaderFooter] = []
        for s in self.sections:
            out.extend(s.all_footers())
        return out


def _iter_para(blocks: list[Block], tables: list[Table]) -> Iterable[Paragraph]:
    by_index = {t.table_index: t for t in tables}
    for b in blocks:
        if isinstance(b, Paragraph):
            yield b
        elif isinstance(b, TableRef):
            table = by_index.get(b.table_index)
            if table is None:
                continue
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_para(cell.blocks, tables)


# ========================================================================
# Field codes (complex w:fldChar and simple w:fldSimple)
# ========================================================================
def classify(instruction: str) -> str:
    """First instruction token, upper-cased (PAGE, NUMPAGES, REF, DATE...)."""
    if not instruction:
        return ""
    tok = instruction.strip().lstrip('"').split()
    return tok[0].upper() if tok else ""


class FieldCollector:
    """Accumulates fields as the paragraph walker streams run tokens."""

    def __init__(self) -> None:
        self.fields: list[Field] = []
        # stack of in-progress complex fields:
        # each entry = [instr_parts, result_parts, phase]  phase: instr|result
        self._stack: list[list] = []

    # -- complex field events -----------------------------------------
    def begin(self) -> None:
        self._stack.append([[], [], "instr"])

    def separate(self) -> None:
        if self._stack:
            self._stack[-1][2] = "result"

    def add_instr(self, text: str) -> None:
        if self._stack and self._stack[-1][2] == "instr":
            self._stack[-1][0].append(text)

    def add_text(self, text: str) -> None:
        """Regular run text; captured as result if a field is open."""
        if self._stack and self._stack[-1][2] == "result":
            self._stack[-1][1].append(text)

    def end(self) -> None:
        if not self._stack:
            return
        instr_parts, result_parts, _ = self._stack.pop()
        instruction = "".join(instr_parts).strip()
        result = "".join(result_parts).strip()
        self.fields.append(Field(
            kind=classify(instruction),
            instruction=instruction,
            result=result,
            simple=False,
        ))

    # -- simple field --------------------------------------------------
    def simple(self, instruction: str, result: str) -> None:
        instruction = (instruction or "").strip()
        self.fields.append(Field(
            kind=classify(instruction),
            instruction=instruction,
            result=(result or "").strip(),
            simple=True,
        ))


# ========================================================================
# Run merging
# ========================================================================
def merge_runs(runs: list[Run]) -> list[Run]:
    merged: list[Run] = []
    for run in runs:
        if run.text == "":
            # keep zero-width control runs from fragmenting merges, but do
            # not emit standalone empty runs
            if merged:
                merged[-1] = Run(text=merged[-1].text, font=merged[-1].font)
            continue
        if merged and merged[-1].font.key() == run.font.key():
            merged[-1] = Run(text=merged[-1].text + run.text, font=merged[-1].font)
        else:
            merged.append(Run(text=run.text, font=run.font))
    return merged


# ========================================================================
# Style / font resolution
# ========================================================================
_HEADING_RE = re.compile(r"^\s*heading\s*([1-9])\s*$", re.IGNORECASE)


def heading_level_from_style(style_id: Optional[str],
                             style_name: Optional[str]) -> Optional[int]:
    for candidate in (style_name, style_id):
        if not candidate:
            continue
        m = _HEADING_RE.match(candidate)
        if m:
            return int(m.group(1))
        # style ids come through as "Heading3" with no space
        m2 = re.match(r"^heading([1-9])$", candidate, re.IGNORECASE)
        if m2:
            return int(m2.group(1))
    return None


_RPR_KEYS = ("font_literal", "font_theme", "size", "bold", "italic",
             "underline", "color")

_PPR_KEYS = ("style_id", "outline_level", "space_before", "space_after",
             "line_spacing", "line_spacing_rule", "alignment",
             "first_line_indent", "left_indent", "right_indent", "is_list")

_ALIGN_MAP = {"both": "JUSTIFY", "distribute": "JUSTIFY", "start": "LEFT",
              "end": "RIGHT", "left": "LEFT", "right": "RIGHT",
              "center": "CENTER"}


def _parse_rpr(rpr: Optional[etree._Element]) -> dict:
    """Parse a w:rPr element into a flat property dict (all keys present,
    None when unset)."""
    d = dict.fromkeys(_RPR_KEYS)
    if rpr is None:
        return d
    rfonts = rpr.find(w("rFonts"))
    if rfonts is not None:
        d["font_literal"] = (rfonts.get(w("ascii"))
                             or rfonts.get(w("hAnsi")))
        d["font_theme"] = (rfonts.get(w("asciiTheme"))
                           or rfonts.get(w("hAnsiTheme")))
    d["size"] = halfpt_to_pt(wval(rpr.find(w("sz"))))
    d["bold"] = as_bool(rpr.find(w("b")))
    d["italic"] = as_bool(rpr.find(w("i")))
    u = rpr.find(w("u"))
    if u is not None:
        d["underline"] = (wval(u) or "single").lower()
    color = rpr.find(w("color"))
    if color is not None:
        val = wval(color)
        theme = color.get(w("themeColor"))
        if val and val != "auto":
            d["color"] = val.upper()
        elif theme:
            d["color"] = f"theme:{theme}"
        elif val:
            d["color"] = val  # "auto"
    return d


def _parse_ppr(ppr: Optional[etree._Element]) -> dict:
    d = dict.fromkeys(_PPR_KEYS)
    if ppr is None:
        return d
    pstyle = ppr.find(w("pStyle"))
    if pstyle is not None:
        d["style_id"] = wval(pstyle)
    outline = ppr.find(w("outlineLvl"))
    if outline is not None:
        try:
            d["outline_level"] = int(wval(outline))
        except (TypeError, ValueError):
            pass
    spacing = ppr.find(w("spacing"))
    if spacing is not None:
        d["space_before"] = twips_to_pt(spacing.get(w("before")))
        d["space_after"] = twips_to_pt(spacing.get(w("after")))
        line = spacing.get(w("line"))
        rule = (spacing.get(w("lineRule")) or "auto").lower()
        if line is not None:
            try:
                line_i = int(line)
            except ValueError:
                line_i = None
            if line_i is not None:
                if rule == "auto":
                    d["line_spacing"] = round(line_i / 240.0, 4)
                    d["line_spacing_rule"] = "AUTO"
                else:
                    d["line_spacing"] = round(line_i / 20.0, 4)
                    d["line_spacing_rule"] = (
                        "EXACT" if rule == "exact" else "AT_LEAST")
    jc = ppr.find(w("jc"))
    if jc is not None:
        d["alignment"] = _map_align(wval(jc))
    ind = ppr.find(w("ind"))
    if ind is not None:
        first = ind.get(w("firstLine"))
        hanging = ind.get(w("hanging"))
        if first is not None:
            d["first_line_indent"] = twips_to_pt(first)
        elif hanging is not None:
            hv = twips_to_pt(hanging)
            d["first_line_indent"] = -hv if hv is not None else None
        left = ind.get(w("left")) or ind.get(w("start"))
        right = ind.get(w("right")) or ind.get(w("end"))
        d["left_indent"] = twips_to_pt(left)
        d["right_indent"] = twips_to_pt(right)
    numpr = ppr.find(w("numPr"))
    if numpr is not None:
        numid = numpr.find(w("numId"))
        if numid is not None and wval(numid) not in (None, "0"):
            d["is_list"] = True
    return d


def _map_align(val: Optional[str]) -> Optional[str]:
    if not val:
        return None
    val = val.lower()
    return _ALIGN_MAP.get(val, val.upper())


class StyleResolver:
    """Owns styles.xml + theme, resolves run/paragraph formatting."""

    def __init__(self, styles_el: Optional[etree._Element],
                 theme_el: Optional[etree._Element]):
        self._styles: dict[str, dict] = {}
        self._default_para: Optional[str] = None
        self._default_char: Optional[str] = None
        self._docdef_rpr: dict = _parse_rpr(None)
        self._docdef_ppr: dict = _parse_ppr(None)
        self.theme_major: Optional[str] = None
        self.theme_minor: Optional[str] = None
        self._style_caches: dict[str, dict] = {"rpr": {}, "ppr": {}}

        self._load_theme(theme_el)
        self._load_styles(styles_el)

    # ---- loading -----------------------------------------------------
    def _load_theme(self, theme_el: Optional[etree._Element]) -> None:
        if theme_el is None:
            return
        def latin(which: str) -> Optional[str]:
            font = theme_el.find(
                f".//{a('fontScheme')}/{a(which)}/{a('latin')}")
            return font.get("typeface") if font is not None else None
        self.theme_major = latin("majorFont")
        self.theme_minor = latin("minorFont")

    def _load_styles(self, styles_el: Optional[etree._Element]) -> None:
        if styles_el is None:
            return
        docdef = styles_el.find(w("docDefaults"))
        if docdef is not None:
            rprd = get(docdef, "rPrDefault", "rPr")
            pprd = get(docdef, "pPrDefault", "pPr")
            self._docdef_rpr = _parse_rpr(rprd)
            self._docdef_ppr = _parse_ppr(pprd)
        for st in styles_el.findall(w("style")):
            sid = st.get(w("styleId"))
            if not sid:
                continue
            stype = st.get(w("type")) or "paragraph"
            is_default = (st.get(w("default")) or "0") in ("1", "true")
            name_el = st.find(w("name"))
            name = wval(name_el)
            based = wval(st.find(w("basedOn")))
            self._styles[sid] = {
                "id": sid,
                "type": stype,
                "name": name,
                "based_on": based,
                "default": is_default,
                "rpr": _parse_rpr(st.find(w("rPr"))),
                "ppr": _parse_ppr(st.find(w("pPr"))),
            }
            if is_default and stype == "paragraph" and not self._default_para:
                self._default_para = sid
            if is_default and stype == "character" and not self._default_char:
                self._default_char = sid

    # ---- style chain merge ------------------------------------------
    def _chain(self, style_id: Optional[str]) -> list[str]:
        """basedOn chain root..leaf for a style id (leaf last)."""
        seen: list[str] = []
        cur = style_id
        guard = 0
        while cur and cur in self._styles and cur not in seen and guard < 50:
            seen.append(cur)
            cur = self._styles[cur]["based_on"]
            guard += 1
        seen.reverse()
        return seen

    def _merged(self, style_id: Optional[str], which: str) -> dict:
        """Flatten a basedOn chain for "rpr" or "ppr" (leaf wins)."""
        keys = _RPR_KEYS if which == "rpr" else _PPR_KEYS
        if not style_id or style_id not in self._styles:
            return dict.fromkeys(keys)
        cache = self._style_caches[which]
        if style_id not in cache:
            merged = dict.fromkeys(keys)
            for sid in self._chain(style_id):
                _overlay(merged, self._styles[sid][which])
            cache[style_id] = merged
        return cache[style_id]

    # ---- public resolution ------------------------------------------
    def resolve_font(self, run_rpr: Optional[etree._Element],
                     para_style_id: Optional[str]) -> ResolvedFont:
        direct = _parse_rpr(run_rpr)
        # character style referenced by the run
        rstyle_el = run_rpr.find(w("rStyle")) if run_rpr is not None else None
        char_style_id = wval(rstyle_el)

        levels = [
            direct,
            self._merged(char_style_id, "rpr"),
            self._merged(para_style_id or self._default_para, "rpr"),
            self._docdef_rpr,
        ]
        name, is_theme = self._resolve_name(levels)
        return ResolvedFont(
            name=name,
            name_is_theme=is_theme,
            size=_first(levels, "size"),
            bold=_first(levels, "bold"),
            italic=_first(levels, "italic"),
            underline=_first(levels, "underline"),
            color=_first(levels, "color"),
        )

    def _resolve_name(self, levels: list[dict]) -> tuple[Optional[str], bool]:
        for lvl in levels:
            theme_tok = lvl.get("font_theme")
            literal = lvl.get("font_literal")
            if theme_tok:
                resolved = self._theme_font(theme_tok)
                if resolved:
                    return resolved, True
            if literal:
                return literal, False
        return None, False

    def _theme_font(self, token: str) -> Optional[str]:
        t = token.lower()
        if t.startswith("major"):
            return self.theme_major
        if t.startswith("minor"):
            return self.theme_minor
        return None

    def resolve_para(self, ppr_el: Optional[etree._Element]) -> ParagraphProps:
        direct = _parse_ppr(ppr_el)
        style_id = direct.get("style_id") or self._default_para
        style_merged = self._merged(style_id, "ppr")
        levels = [direct, style_merged, self._docdef_ppr]
        style_name = None
        if style_id and style_id in self._styles:
            style_name = self._styles[style_id]["name"]
        heading = heading_level_from_style(style_id, style_name)
        outline = _first(levels, "outline_level")
        if outline is None and heading is not None:
            outline = heading - 1
        return ParagraphProps(
            style_id=style_id,
            style_name=style_name,
            outline_level=outline,
            heading_level=heading,
            space_before=_first(levels, "space_before"),
            space_after=_first(levels, "space_after"),
            line_spacing=_first(levels, "line_spacing"),
            line_spacing_rule=_first(levels, "line_spacing_rule"),
            alignment=_first(levels, "alignment"),
            first_line_indent=_first(levels, "first_line_indent"),
            left_indent=_first(levels, "left_indent"),
            right_indent=_first(levels, "right_indent"),
            is_list=bool(_first(levels, "is_list")),
        )

    # ---- public summary ---------------------------------------------
    def index(self) -> StyleIndex:
        styles = {
            sid: StyleInfo(
                style_id=sid,
                name=info["name"],
                type=info["type"],
                based_on=info["based_on"],
                is_default=info["default"],
            )
            for sid, info in self._styles.items()
        }
        return StyleIndex(
            styles=styles,
            theme_major=self.theme_major,
            theme_minor=self.theme_minor,
            default_para_style=self._default_para,
            default_char_style=self._default_char,
        )


def _overlay(base: dict, over: dict) -> None:
    """Apply non-None values of ``over`` onto ``base`` in place."""
    for k, v in over.items():
        if v is not None:
            base[k] = v


def _first(levels: list[dict], key: str):
    for lvl in levels:
        v = lvl.get(key)
        if v is not None:
            return v
    return None


# ========================================================================
# Sections, headers and footers
# ========================================================================
_HEADER_VARIANTS = (
    ("default", "header"),
    ("first", "first_page_header"),
    ("even", "even_page_header"),
)
_FOOTER_VARIANTS = (
    ("default", "footer"),
    ("first", "first_page_footer"),
    ("even", "even_page_footer"),
)


def build_sections(document, resolver) -> list[Section]:
    sections: list[Section] = []
    prev: Optional[Section] = None
    for i, sec in enumerate(document.sections):
        section = Section(
            index=i,
            different_first_page=bool(sec.different_first_page_header_footer),
        )
        for which, attr in _HEADER_VARIANTS:
            hf = _capture(getattr(sec, attr, None), which, "header",
                          i, resolver, prev)
            if hf is not None:
                section.headers[which] = hf
        for which, attr in _FOOTER_VARIANTS:
            hf = _capture(getattr(sec, attr, None), which, "footer",
                          i, resolver, prev)
            if hf is not None:
                section.footers[which] = hf
        sections.append(section)
        prev = section
    return sections


def _capture(hf_obj, which: str, kind: str, index: int, resolver,
             prev: Optional[Section]) -> Optional[HeaderFooter]:
    if hf_obj is None:
        return None

    linked = bool(getattr(hf_obj, "is_linked_to_previous", False))
    if linked:
        base = _inherit(prev, which, kind)
        if base is None:
            return None
        return HeaderFooter(
            which=which, kind=kind, section_index=index,
            paragraphs=base.paragraphs, fields=base.fields,
            is_linked_to_previous=True,
        )

    root = _root_element(hf_obj)
    if root is None:
        return None

    # Ctx / walk_container are defined further down in this module; the
    # former lazy import existed only to break a document<->sections cycle.
    ctx = Ctx(resolver)
    blocks = walk_container(root, ctx)
    paras: list[Paragraph] = [b for b in blocks if isinstance(b, Paragraph)]
    for table in ctx.tables:
        paras.extend(table.iter_paragraphs())
    fields = [f for p in paras for f in p.fields]

    if not paras and not fields:
        return None
    return HeaderFooter(
        which=which, kind=kind, section_index=index,
        paragraphs=paras, fields=fields, is_linked_to_previous=False,
    )


def _inherit(prev: Optional[Section], which: str, kind: str):
    if prev is None:
        return None
    table = prev.headers if kind == "header" else prev.footers
    return table.get(which) or table.get("default")


def _root_element(hf_obj) -> Optional[etree._Element]:
    """Best-effort access to the header/footer root (CT_HdrFtr) element."""
    el = getattr(hf_obj, "_element", None)
    if el is not None:
        return el
    # fall back through the first paragraph's parent element
    try:
        paras = hf_obj.paragraphs
    except Exception:
        paras = []
    if paras:
        p = getattr(paras[0], "_p", None) or getattr(paras[0], "_element", None)
        if p is not None:
            return p.getparent()
    return None


# ========================================================================
# Document walk -- the entry point
# ========================================================================
THEME_RELTYPE = ("http://schemas.openxmlformats.org/officeDocument/2006/"
                 "relationships/theme")


class Ctx:
    """Mutable walk state: the global block counter and the table registry."""

    def __init__(self, resolver: StyleResolver):
        self.resolver = resolver
        self.tables: list[Table] = []
        self._counter = -1

    def next_index(self) -> int:
        self._counter += 1
        return self._counter


# --------------------------------------------------------------------------
# Public entry point
# --------------------------------------------------------------------------
# ========================================================================
# Inferred headings
# ========================================================================
# Plenty of real SOPs never apply Word's heading styles: the author just
# bolds a line and bumps its size. Word records nothing for those, so
# outline_level and heading_level are both None and the document looks to
# the rules like one long unstructured run of text. This pass recovers the
# visual structure -- but only for documents that declared no heading
# anywhere, so a properly styled file is always taken at its word.
_SENTENCE_END = re.compile(r"[.;,!?]\s*$")
MAX_HEADING_WORDS = 12
MIN_PARAS_TO_INFER = 4
# bold only means "heading" while it stays the exception; in a document set
# mostly in bold it carries no signal at all
BOLD_SIGNAL_MAX_SHARE = 0.4


def _char_weighted_size(paras: list["Paragraph"]) -> Optional[float]:
    counter: Counter = Counter()
    for p in paras:
        for r in p.runs:
            if r.text.strip() and r.font.size is not None:
                counter[r.font.size] += len(r.text)
    return counter.most_common(1)[0][0] if counter else None


def _para_size(p: "Paragraph") -> Optional[float]:
    return _char_weighted_size([p])


def _para_is_bold(p: "Paragraph") -> bool:
    """True when the paragraph's visible text is predominantly bold."""
    bold = plain = 0
    for r in p.runs:
        n = len(r.text.strip())
        if not n:
            continue
        if r.font.bold:
            bold += n
        else:
            plain += n
    return bold > plain


def infer_heading_levels(blocks: list[Block], tables: list[Table]) -> None:
    """Fill ``props.inferred_heading_level`` for paragraphs that read as
    headings: short, not a sentence, and either larger than the body text or
    bold where bold is rare. Levels come from ranking the sizes found, so the
    biggest line becomes 1 and the next distinct size 2."""
    everything = list(_iter_para(blocks, tables))
    if any(p.props.heading_level is not None or p.props.outline_level is not None
           for p in everything):
        return                      # Word declared its own structure

    body = [p for p in everything
            if not p.in_table and not p.from_textbox and p.text.strip()]
    if len(body) < MIN_PARAS_TO_INFER:
        return

    body_size = _char_weighted_size(body)
    bold_share = sum(1 for p in body if _para_is_bold(p)) / len(body)
    bold_is_signal = bold_share <= BOLD_SIGNAL_MAX_SHARE

    candidates: list[tuple[Paragraph, Optional[float]]] = []
    for p in body:
        text = p.text.strip()
        if p.props.is_list or len(text.split()) > MAX_HEADING_WORDS:
            continue
        if _SENTENCE_END.search(text):
            continue
        size = _para_size(p)
        larger = (size is not None and body_size is not None
                  and size > body_size)
        if larger or (bold_is_signal and _para_is_bold(p)):
            candidates.append((p, size))
    if not candidates:
        return

    ranked = sorted({s for _, s in candidates if s is not None}, reverse=True)
    for para, size in candidates:
        level = ranked.index(size) + 1 if size in ranked else len(ranked) + 1
        para.props.inferred_heading_level = min(level, 9)


def build_doc(file, filename: str) -> Doc:
    document = OpenDocument(file)
    resolver = StyleResolver(_styles_element(document), _theme_element(document))

    ctx = Ctx(resolver)
    body = document.element.body
    blocks = walk_container(body, ctx)

    infer_heading_levels(blocks, ctx.tables)

    sections = build_sections(document, resolver)
    core = _core_props(document)

    return Doc(
        filename=filename,
        core=core,
        blocks=blocks,
        tables=ctx.tables,
        sections=sections,
        styles=resolver.index(),
    )


# --------------------------------------------------------------------------
# Block-level walk
# --------------------------------------------------------------------------
def walk_container(container: etree._Element, ctx: Ctx,
                   in_table: bool = False,
                   table_pos: Optional[tuple] = None,
                   from_textbox: bool = False) -> list[Block]:
    """Walk one container (body, table cell, sdtContent, header/footer)."""
    blocks: list[Block] = []
    for child in container:
        tag = local(child)
        if tag == "p":
            idx = ctx.next_index()
            para, textboxes = parse_paragraph(
                child, ctx.resolver, idx, in_table, table_pos, from_textbox)
            blocks.append(para)
            for tb in textboxes:
                blocks.extend(
                    walk_container(tb, ctx, in_table=in_table,
                                   table_pos=table_pos, from_textbox=True))
        elif tag == "tbl":
            blocks.append(_walk_table(child, ctx))
        elif tag == "sdt":
            content = child.find(w("sdtContent"))
            if content is not None:
                blocks.extend(walk_container(
                    content, ctx, in_table=in_table, table_pos=table_pos,
                    from_textbox=from_textbox))
        # sectPr, bookmarks, proofErr, etc. carry no block content
    return blocks


def _walk_table(tbl: etree._Element, ctx: Ctx) -> TableRef:
    idx = ctx.next_index()
    table_index = len(ctx.tables)
    table = Table(table_index=table_index, block_index=idx, rows=[])
    ctx.tables.append(table)  # reserve slot before recursing into cells

    rows: list[Row] = []
    for r, tr in enumerate(tbl.findall(w("tr"))):
        cells: list[Cell] = []
        for c, tc in enumerate(tr.findall(w("tc"))):
            cell_blocks = walk_container(
                tc, ctx, in_table=True, table_pos=(table_index, r, c))
            cells.append(Cell(blocks=cell_blocks))
        rows.append(Row(cells=cells))
    table.rows = rows
    return TableRef(block_index=idx, table_index=table_index,
                    location=f"Table {table_index + 1}")


# --------------------------------------------------------------------------
# Paragraph parsing
# --------------------------------------------------------------------------
def parse_paragraph(p: etree._Element, resolver: StyleResolver,
                    block_index: int, in_table: bool,
                    table_pos: Optional[tuple],
                    from_textbox: bool) -> tuple[Paragraph, list]:
    props = resolver.resolve_para(p.find(w("pPr")))
    props.in_table = in_table

    sink = _RunSink(resolver, props.style_id)
    _walk_runs(p, sink)

    para = Paragraph(
        block_index=block_index,
        text="".join(sink.text),
        runs=merge_runs(sink.runs),
        props=props,
        fields=sink.fields.fields,
        in_table=in_table,
        from_textbox=from_textbox,
        table_pos=table_pos,
        location=_para_location(block_index, in_table, table_pos, from_textbox),
    )
    return para, sink.textboxes


_RECURSE_TAGS = {"hyperlink", "ins", "smartTag", "customXml"}


class _RunSink:
    """Accumulator for one paragraph: runs, text, fields and text boxes.

    Threading these six as arguments meant a seven-parameter signature
    restated at every recursion site.
    """

    __slots__ = ("resolver", "style_id", "fields", "runs", "text", "textboxes")

    def __init__(self, resolver: StyleResolver, style_id: Optional[str]):
        self.resolver = resolver
        self.style_id = style_id
        self.fields = FieldCollector()
        self.runs: list[Run] = []
        self.text: list[str] = []
        self.textboxes: list[etree._Element] = []

    def emit(self, ch: str, font: ResolvedFont) -> None:
        self.runs.append(Run(text=ch, font=font))
        self.text.append(ch)
        self.fields.add_text(ch)


def _walk_runs(elem: etree._Element, sink: _RunSink) -> None:
    for child in elem:
        tag = local(child)
        if tag == "r":
            _walk_single_run(child, sink)
        elif tag == "fldSimple":
            instr = child.get(w("instr")) or ""
            before = len(sink.text)
            _walk_runs(child, sink)
            sink.fields.simple(instr, "".join(sink.text[before:]))
        elif tag == "sdt":
            content = child.find(w("sdtContent"))
            if content is not None:
                _walk_runs(content, sink)
        elif tag in _RECURSE_TAGS:
            _walk_runs(child, sink)
        elif tag == "del":
            continue  # tracked deletion: not part of the accepted document


def _walk_single_run(run: etree._Element, sink: _RunSink) -> None:
    font = sink.resolver.resolve_font(run.find(w("rPr")), sink.style_id)
    for rc in run:
        tag = local(rc)
        if tag == "t":
            sink.emit(rc.text or "", font)
        elif tag == "instrText":
            sink.fields.add_instr(rc.text or "")
        elif tag == "delInstrText":
            continue
        elif tag == "fldChar":
            ftype = rc.get(w("fldCharType"))
            if ftype == "begin":
                sink.fields.begin()
            elif ftype == "separate":
                sink.fields.separate()
            elif ftype == "end":
                sink.fields.end()
        elif tag == "tab":
            sink.emit("\t", font)
        elif tag in ("br", "cr"):
            sink.emit("\n", font)
        elif tag == "noBreakHyphen":
            sink.emit("-", font)
        elif tag in ("drawing", "pict", "object"):
            for txbx in rc.iter(w("txbxContent")):
                sink.textboxes.append(txbx)


def _para_location(idx: int, in_table: bool, table_pos: Optional[tuple],
                   from_textbox: bool) -> str:
    if in_table and table_pos is not None:
        t, r, c = table_pos
        return f"Table {t + 1}, row {r + 1}, cell {c + 1}"
    if from_textbox:
        return f"Text box (block {idx})"
    return f"Paragraph {idx}"


# --------------------------------------------------------------------------
# Part helpers
# --------------------------------------------------------------------------
def _styles_element(document) -> Optional[etree._Element]:
    try:
        return document.styles.element
    except Exception:
        return None


def _theme_element(document) -> Optional[etree._Element]:
    try:
        for rel in document.part.rels.values():
            if rel.reltype == THEME_RELTYPE:
                try:
                    return etree.fromstring(rel.target_part.blob)
                except Exception:
                    return None
    except Exception:
        return None
    return None


def _core_props(document) -> CoreProps:
    cp = document.core_properties
    return CoreProps(
        title=cp.title or None,
        author=cp.author or None,
        subject=cp.subject or None,
        keywords=cp.keywords or None,
        created=cp.created,
        modified=cp.modified,
        last_modified_by=cp.last_modified_by or None,
    )
